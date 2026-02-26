#!/usr/bin/env python3
"""
Генератор синопсиса протокола биоэквивалентности.

Заполняет шаблон «Шаблон Синопсиса Протокола.docx» данными из JSON.

Использование:
    python generate_synopsis.py [input.json] [output.docx]

По умолчанию:
    python generate_synopsis.py recommendation_amlodipine.json synopsis_output.docx
"""

import argparse
import json
import math
import sys
from pathlib import Path
from docx import Document

# ─── Константы ────────────────────────────────────────────────────────────────

PLACEHOLDER = "\u2002" * 5   # 5 неразрывных полупробелов — заполнитель в шаблоне

# ─── Дизайн: падежная форма для заголовка (строка 0) ────────────────────────
# Ключи — то, что приходит из LLM в поле "design".

DESIGN_TITLE_DESCRIPTIONS: dict[str, str] = {
    "2×2 Cross-over":        "двухпериодном двухпоследовательном перекрёстном",
    "3-way Replicate":       "трёхпериодном реплицированном перекрёстном",
    "4-way Replicate":       "четырёхпериодном полностью реплицированном перекрёстном",
    "Параллельный":          "параллельно-групповом",
    # русские варианты (legacy)
    "2×2 Перекрёстный":     "двухпериодном двухпоследовательном перекрёстном",
    "3-way Реплицированный": "трёхпериодном реплицированном перекрёстном",
    "4-way Реплицированный": "четырёхпериодном полностью реплицированном перекрёстном",
}

# ─── Стандартные критерии отбора здоровых добровольцев ───────────────────────
# Источник: примеры синопсисов (Тенофовир, Биктегравир).
# {drug} подставляется автоматически; <...> — ручное заполнение.

INCLUSION_CRITERIA = (
    "1. Мужчины и/или женщины в возрасте от <18> до <45> лет включительно.\n"
    "2. Индекс массы тела (ИМТ) 18,5–30,0 кг/м²; масса тела >45 кг (женщины), >55 кг (мужчины), ≤110 кг.\n"
    "3. Артериальное давление: систолическое 100–129 мм рт.ст., диастолическое 60–89 мм рт.ст.\n"
    "4. ЧСС 60–89 уд/мин; ЧДД 12–20 в мин; температура тела 36,0–36,9°C.\n"
    "5. Верифицированный диагноз «практически здоров» по результатам скрининга.\n"
    "6. Отрицательные результаты тестов на ВИЧ, гепатиты B и C, сифилис.\n"
    "7. Отрицательные тесты на наркотические вещества, алкоголь, котинин при скрининге.\n"
    "8. Способность и готовность выполнять требования протокола; подписанное информированное согласие.\n"
    "9. Согласие на применение надёжных методов контрацепции в течение исследования и <1> месяца после окончания.\n"
    "10. Готовность воздерживаться от алкоголя не менее 72 ч до каждого визита."
)

NON_INCLUSION_CRITERIA = (
    "1. Хронические и/или острые заболевания сердечно-сосудистой, дыхательной, нервной, эндокринной, "
    "опорно-двигательной, кроветворной, иммунной систем, почек, печени, ЖКТ, кожи; онкологические заболевания.\n"
    "2. Хирургические вмешательства на органах ЖКТ (кроме аппендэктомии ≥1 года назад).\n"
    "3. Острые инфекционные заболевания в течение 4 недель до скрининга.\n"
    "4. Приём любых ЛС в течение 4 недель до начала исследования. "
    "<Дополнить: специфические группы препаратов, влияющих на ФК {drug}.>\n"
    "5. Донация крови или её компонентов в течение 2 месяцев до скрининга.\n"
    "6. Приём гормональных контрацептивов в течение 2 месяцев до начала исследования.\n"
    "7. Клинически значимые отклонения лабораторных показателей (ОАК, биохимия, ОАМ) при скрининге.\n"
    "8. Положительные тесты на алкоголь, наркотики, беременность, котинин, гепатиты B/C, сифилис, ВИЧ.\n"
    "9. Гиперчувствительность к {drug} и/или вспомогательным веществам в анамнезе.\n"
    "10. <Непереносимость лактозы — указать если применимо.>\n"
    "11. Употребление алкоголя >10 единиц в неделю; алкогольная зависимость.\n"
    "12. <Особые диетические ограничения — грейпфрут, кофеин и пр. — уточнить при необходимости.>\n"
    "13. Беременность, лактация (для женщин).\n"
    "14. Участие в другом клиническом исследовании в течение 3 месяцев до начала настоящего."
)

EXCLUSION_CRITERIA = (
    "1. Добровольный отзыв информированного согласия.\n"
    "2. Грубое нарушение протокола исследования.\n"
    "3. Нарушение критериев включения/невключения, выявленное после рандомизации.\n"
    "4. Угроза безопасности добровольца по оценке исследователя.\n"
    "5. Развитие НЯ или СНЯ, требующего прекращения участия.\n"
    "6. Необходимость применения ЛС, способных влиять на ФК {drug}.\n"
    "7. Рвота и/или диарея в течение <2×Tmax — указать значение в часах> после приёма дозы.\n"
    "8. Положительные повторные тесты на наркотики, алкоголь, котинин или беременность.\n"
    "9. Иные обстоятельства, по мнению исследователя препятствующие участию добровольца."
)

# ─── Шаблонные тексты по блокам синопсиса ────────────────────────────────────
# Используй {Drug} для названия с заглавной буквы, {drug} — строчными.
# Режим __REPLACE__ заменяет ВСЁ содержимое ячейки этим текстом.
DOSING_DESCRIPTIONS: dict[str, str] = {
    "однократный": ["однократного", "однократном"],
    "многократный": ["многократного", "многократном"],
}

SECTION_TEMPLATES: dict[str, str] = {

    # Строка 8 — Цель исследования
    "study_goal1": (
        "Основная цель:\n"
        "Оценка сравнительной фармакокинетики и биоэквивалентности препаратов («Название тестового препарата») и {Drug} («Название референтного препарата») {fed_state} у здоровых добровольцев."
    ),
    "study_goal2": (
        "\nДополнительная цель: \n"
        "Сравнительная оценка безопасности {dosing_description} приема препаратов («[ИНН] тестового препарата») и {Drug} («[ИНН] референтного препарата») у здоровых добровольцев."
    ),

    # Строка 9 — Задачи исследования
    "study_tasks": (
        "1. Определить концентрацию {drug} в плазме крови добровольцев после "
        "{dosing1} приёма препарата сравнимых препаратов («[ИНН] тестового препарата») и {Drug} («[ИНН] референтного препарата»).\n"
        "2. Оценить фармакокинетические параметры и относительную биодоступность сравниваемых препаратов.\n"
        "3. Оценить биоэквивалентность сравниваемых препаратов на основе статистического анализа фармакокинетических данных.\n"
        "4. Оценить профиль безопасности сравниваемых препаратов при {dosing2} применении "
        "частоту возникновения нежелательных явлений (НЯ)/ "
        "серьезных нежелательных явлений (СНЯ), изменения данных лабораторных "
        "исследований, физикального осмотра, функций жизненно важных органов, "
        "показателей электрокардиограммы (ЭКГ))."
    ),

    # Строка 12 — Количество добровольцев / обоснование расчёта
    "sample_size_justification": (
        "Количество добровольцев рассчитано на основании данных о внутрисубъектной "
        "вариабельности (CVintra = {cv_intra}%) для {drug} (Cmax и/или AUC) при уровне "
        "значимости α = 0,05, мощности 80% и ожидаемом выбывании 20%.\n"
        "Базовый размер выборки: {n_total} участников.\n"
        "С учётом выбывания: {n_subjects} участников."
    ),

    # Строка 20 — Фармакокинетические параметры
    "pk_parameters": (
        "Фармакокинетика {drug} будет оценена по следующим параметрам:\n"
        "• Cmax — максимальная концентрация {drug} в плазме крови;\n"
        "• AUC(0-t) — площадь под фармакокинетической кривой «концентрация–время» "
        "от 0 до последней измеримой концентрации {drug};\n"
        "• AUC(0-∞) — AUC, экстраполированная до бесконечности;\n"
        "• Tmax — время достижения Cmax;\n"
        "• T½ — период полувыведения."
    ),

    # Строка 22 — Критерии биоэквивалентности
    "be_criteria": (
        "Препарат {Drug} считается биоэквивалентным референтному препарату, если 90% "
        "доверительные интервалы отношений геометрических средних тест/референт для Cmax "
        "и AUC(0-t) находятся в пределах 80,00–125,00% в соответствии с действующими "
        "требованиями ЕАЭС."
    ),

    # Строка 25 — Методы статистического анализа
    "statistical_methods": (
        "Статистический анализ биоэквивалентности препарата {Drug} проводится методом "
        "дисперсионного анализа (ANOVA) в логарифмической шкале. Рассчитываются точечные "
        "оценки и 90% доверительные интервалы (ДИ) отношений геометрических средних "
        "тест/референт для Cmax и AUC(0-t). Биоэквивалентность устанавливается, если "
        "90% ДИ для обоих параметров находятся в диапазоне 80,00–125,00%."
    ),
}


def _render(template_key: str, **kwargs) -> str:
    """Форматирует шаблонный текст блока синопсиса, подставляя параметры."""
    return SECTION_TEMPLATES[template_key].format(**kwargs)


# ─── Расчёт хронологии исследования ──────────────────────────────────────────

def _days_word(n: int) -> str:
    """Возвращает правильную форму слова 'день' для числа n."""
    last2, last1 = n % 100, n % 10
    if 11 <= last2 <= 14:
        return "дней"
    if last1 == 1:
        return "день"
    if 2 <= last1 <= 4:
        return "дня"
    return "дней"


def _compute_study_timeline(data: dict) -> dict:
    """
    Вычисляет хронологические параметры исследования на основе данных из JSON.

    Args:
        data: словарь из JSON-файла рекомендации (содержит design, t_half_used, pk_data, n_subjects)

    Returns:
        Словарь с вычисленными параметрами хронологии.
    """
    design    = data.get("design", "2×2 Cross-over")
    t_half    = data.get("t_half_used") or 24.0

    # Tmax из первой записи pk_data (обычно merged)
    tmax = None
    for r in data.get("pk_data", []):
        if r.get("tmax"):
            tmax = r["tmax"]
            break
    tmax = tmax or 2.0

    # ── Длительность забора ПК-образцов ──────────────────────────────────────
    # Минимум 24 ч, максимум 72 ч (практическое ограничение), не менее 3×T½
    pk_sampling_h    = int(min(72, max(24, math.ceil(5 * t_half))))
    pk_sampling_days = math.ceil(pk_sampling_h / 24)

    # ── Длительность периода (дней в центре) ─────────────────────────────────
    # День 0 = госпитализация, День 1 = приём, День pk_sampling_days = выписка
    period_days = pk_sampling_days + 1

    # ── Отмывочный период: ≥5×T½, минимум 7 дней ─────────────────────────────
    washout_days = max(7, math.ceil(5 * t_half / 24))

    # ── Период последующего наблюдения ────────────────────────────────────────
    followup_day = 7

    # ── Количество периодов ───────────────────────────────────────────────────
    n_periods_map = {
        "2×2 Cross-over":  2,
        "3-way Replicate": 3,
        "4-way Replicate": 4,
        "Параллельный":    1,
    }
    n_periods  = n_periods_map.get(design, 2)
    n_washouts = max(0, n_periods - 1)

    n_subjects = data.get("n_subjects") or 12

    # ── Суммарная длительность ───────────────────────────────────────────────
    if n_periods == 1:  # параллельный
        total_days = 14 + period_days + followup_day
    else:
        total_days = 14 + period_days * n_periods + washout_days * n_washouts + followup_day

    # ── Дни приёма дозы (1, 1+washout, 1+2*washout, ...) ────────────────────
    period_dosing_days   = [1 + i * washout_days for i in range(n_periods)]
    period_admission_days = [d - 1 for d in period_dosing_days]   # День 0, День p2-1, ...
    # Выписка: dosing_day + pk_sampling_days - 1
    #   Для pk_sampling_days=1: discharge = dosing_day (т.е. выписка в день приёма)
    #   Для pk_sampling_days=3: discharge = dosing_day + 2
    period_discharge_days = [d + pk_sampling_days - 1 for d in period_dosing_days]

    last_dosing_day      = period_dosing_days[-1]
    last_discharge_day   = period_discharge_days[-1]
    followup_visit_day   = last_dosing_day + followup_day

    # ── Стандартный ПК-план: число точек ─────────────────────────────────────
    # Приблизительно: 8 базовых + 2 точки на каждые 8 ч забора
    n_timepoints = min(20, max(10, 8 + pk_sampling_days * 2))

    # ── Объёмы крови ─────────────────────────────────────────────────────────
    blood_per_sample    = 5    # мл на ПК-образец
    dead_vol_per_draw   = 0.5  # мл мёртвого объёма на забор
    total_pk_samples    = n_timepoints * n_periods   # на одного добровольца
    pk_blood_volume     = blood_per_sample * total_pk_samples
    dead_vol_total      = round(dead_vol_per_draw * total_pk_samples, 1)
    lab_blood           = 40   # мл на клинические, биохимические анализы (стандарт)
    total_blood_volume  = pk_blood_volume + int(dead_vol_total) + lab_blood
    total_samples_to_lab = total_pk_samples * n_subjects

    return {
        "design":               design,
        "n_periods":            n_periods,
        "screening_days":       14,
        "period_days":          period_days,
        "pk_sampling_h":        pk_sampling_h,
        "pk_sampling_days":     pk_sampling_days,
        "washout_days":         washout_days,
        "followup_day":         followup_day,
        "total_days":           total_days,
        "n_subjects":           n_subjects,
        "period_dosing_days":   period_dosing_days,
        "period_admission_days":period_admission_days,
        "period_discharge_days":period_discharge_days,
        "last_dosing_day":      last_dosing_day,
        "last_discharge_day":   last_discharge_day,
        "followup_visit_day":   followup_visit_day,
        "n_timepoints":         n_timepoints,
        "total_pk_samples":     total_pk_samples,
        "pk_blood_volume":      pk_blood_volume,
        "dead_vol_total":       dead_vol_total,
        "dead_vol_per_draw":    dead_vol_per_draw,
        "lab_blood":            lab_blood,
        "total_blood_volume":   total_blood_volume,
        "total_samples_to_lab": total_samples_to_lab,
        "t_half":               t_half,
    }


def _generate_methodology_text(tl: dict, drug: str, dosing: str, fed_state: str) -> str:
    """Генерирует текст методологии исследования для строки 11."""
    design   = tl["design"]
    n_periods = tl["n_periods"]

    pd  = tl["period_days"]
    wo  = tl["washout_days"]
    fup = tl["followup_day"]
    psh = tl["pk_sampling_h"]
    dos = tl["period_dosing_days"]
    dis = tl["period_discharge_days"]
    fv  = tl["followup_visit_day"]
    t_h = tl["t_half"]
    n_pks = tl["total_pk_samples"]

    drug_cap = drug.capitalize()
    dosing_adv = "однократного" if dosing == "однократный" else "многократного"

    if design == "Параллельный":
        group_descr = (
            "Добровольцы будут рандомизированы в одну из двух групп в соотношении 1:1: "
            "группа исследуемого препарата и группа референтного препарата."
        )
        period_descr = (
            f"Исследование включает один период ФК исследования.\n"
            f"Длительность периода скрининга составит не более 14 дней, "
            f"длительность периода ФК исследования составит {pd} дней, "
            f"периода последующего наблюдения – {fup} дней от приёма препарата.\n\n"
            f"Каждый доброволец получит однократно исследуемый или референтный препарат в День 1 {fed_state}, "
            f"запивая 200 мл ±10 мл бутилированной негазированной воды. "
            f"Добровольцы останутся в центре в течение как минимум {psh} часов после дозирования "
            f"(до Дня {dis[0]}) с целью отбора биообразцов для фармакокинетического анализа.\n\n"
            f"Период последующего наблюдения (День {fv}) проводится через {fup} дней после приёма препарата."
        )
    else:
        # Перекрёстный или реплицированный дизайн
        period_names_ru = {2: "двух", 3: "трёх", 4: "четырёх"}
        n_str = period_names_ru.get(n_periods, str(n_periods))
        group_descr = (
            f"Добровольцы будут распределены в соответствии с рандомизационным списком "
            f"в одну из групп в соотношении 1:1."
        )
        period_descr = (
            f"Исследование будет состоять из периода скрининга, {n_str} периодов ФК исследования "
            f"с отмывочными периодами между ними и периода последующего наблюдения.\n"
            f"Длительность периода скрининга составит не более 14 дней, "
            f"длительность каждого периода ФК исследования составит {pd} дней, "
            f"длительность отмывочного периода – {wo} дней от приёма препарата в предшествующем периоде, "
            f"периода последующего наблюдения – {fup} дней от приёма препарата в последнем периоде.\n\n"
        )

        # Описание периодов
        dosing_days_str = " и ".join(f"День {d}" for d in dos)
        period_descr += (
            f"Добровольцы будут госпитализированы вечером накануне приёма препаратов. "
            f"Утром в {dosing_days_str} добровольцы получат {dosing_adv} дозу "
            f"исследуемого/референтного препарата {fed_state}, "
            f"запивая 200 мл ±10 мл бутилированной негазированной воды комнатной температуры.\n\n"
            f"Добровольцы останутся в центре в течение как минимум {psh} часов после дозирования "
            f"с целью отбора биообразцов для анализа фармакокинетики и оценки параметров безопасности.\n\n"
            f"В каждом из {n_periods} периодов ФК исследования будет проводиться отбор "
            f"{tl['n_timepoints']} проб крови по 5 мл у каждого добровольца. "
            f"Всего в ходе исследования будет отобрано {n_pks} проб крови на одного добровольца "
            f"для оценки фармакокинетических параметров {drug}.\n\n"
        )

        # Отмывочный период
        period_descr += (
            f"Отмывочный период составит {wo} дней с момента приёма дозы. "
            f"Длительность отмывочного периода превышает длительность 5-ти периодов полувыведения {drug_cap} "
            f"(Т½ {drug_cap} из плазмы крови составляет около {t_h} ч).\n\n"
        )

        # Период наблюдения
        period_descr += (
            f"Визит периода последующего наблюдения будет проведён на {fup} день "
            f"(День {fv}) от последнего приёма препарата с целью оценки НЯ/СНЯ."
        )

    return (
        "Настоящее исследование будет выполнено с участием здоровых добровольцев, "
        "соответствующих критериям включения/невключения и подписавших «Информационный листок "
        "добровольца с формой информированного согласия».\n\n"
        f"{group_descr}\n\n"
        f"{period_descr}"
    )


def _generate_periods_text(tl: dict, drug: str, dosing: str, fed_state: str) -> str:
    """Генерирует текст для строки 18 — «Периоды исследования»."""
    design   = tl["design"]
    n_periods = tl["n_periods"]
    psd      = tl["pk_sampling_days"]
    wo       = tl["washout_days"]
    fup      = tl["followup_day"]
    dos      = tl["period_dosing_days"]
    adm      = tl["period_admission_days"]
    dis      = tl["period_discharge_days"]
    fv       = tl["followup_visit_day"]

    lines = ["Исследование включает:", ""]
    lines.append("Период скрининга (предварительное обследование добровольцев):")
    lines.append("Визит 0. (День -14 – День -1).")
    lines.append("")

    if design == "Параллельный":
        lines.append("Период ФК исследования:")
        lines.append(f"Визит 1. День {adm[0]} – День {dis[0]} (госпитализация)")
        lines.append(f"Госпитализация и рандомизация – День {adm[0]}")
        lines.append(f"Прием препарата – День {dos[0]}")
        if psd == 1:
            lines.append(f"Отбор образцов крови – День {dos[0]}")
        else:
            lines.append(f"Отбор образцов крови – День {dos[0]} – День {dis[0]}")
        lines.append(f"Завершение госпитализации – День {dis[0]}.")
        lines.append("")
        lines.append("Период последующего наблюдения:")
        lines.append(f"Визит 2. День {fv} (окно визита +2 дня)")
    else:
        for i in range(n_periods):
            v = i + 1
            lines.append(f"Период {v} ФК исследования.")
            lines.append(f"Визит {v}. День {adm[i]} – День {dis[i]} (госпитализация)")
            if i == 0:
                lines.append(f"Госпитализация и рандомизация – День {adm[i]}")
            else:
                lines.append(f"Госпитализация – День {adm[i]}")
            lines.append(f"Прием препарата – День {dos[i]}")
            if psd == 1:
                lines.append(f"Отбор образцов крови для анализа фармакокинетики и оценка параметров безопасности – День {dos[i]}")
            else:
                lines.append(f"Отбор образцов крови для анализа фармакокинетики и оценка параметров безопасности – День {dos[i]} – День {dis[i]}")
            lines.append(f"Завершение госпитализации – День {dis[i]}.")

            if i < n_periods - 1:
                wash_end = dos[i + 1]
                lines.append(f"Отмывочный период: День {dos[i]} – День {wash_end} ({wo} дней от приема препарата в Периоде {v} ФК исследования).")
            lines.append("")

        lines.append("Период последующего наблюдения:")
        lines.append(f"Визит {n_periods + 1}. День {fv} (окно визита +2 дня)")
        lines.append(
            "Доброволец посетит центр через 7 дней с момента последнего приема препарата, "
            "будет осуществлен сбор данных о состоянии добровольца."
        )

    lines += [
        "",
        "Незапланированный визит",
        "Проводится при необходимости. При наличии показаний может быть дополнительно выполнена "
        "любая из процедур исследования по решению Исследователя.",
        "",
        "Визит досрочного завершения участия в исследовании",
        "Проводится при досрочном выбывании добровольца из исследования.",
    ]

    return "\n".join(lines)


def _generate_duration_text(tl: dict) -> str:
    """Генерирует текст для строки 19 — «Продолжительность исследования»."""
    design    = tl["design"]
    n_periods = tl["n_periods"]
    total     = tl["total_days"]
    pd        = tl["period_days"]
    wo        = tl["washout_days"]
    fup       = tl["followup_day"]
    last_dos  = tl["last_dosing_day"]
    last_dis  = tl["last_discharge_day"]
    fv        = tl["followup_visit_day"]

    if design == "Параллельный":
        return (
            f"Максимальная продолжительность участия в исследовании для одного добровольца составит "
            f"{total} {_days_word(total)}. Период скрининга продлится от 1 до 14 дней. "
            f"Длительность периода ФК исследования составит {pd} {_days_word(pd)} "
            f"(включая госпитализацию вечером накануне приёма препарата). "
            f"Период наблюдения – {fup} {_days_word(fup)} (День {last_dos} – День {fv}) от приёма дозы."
        )

    # Перекрёстный / реплицированный
    period_names = {2: "1 и 2", 3: "1, 2 и 3", 4: "1, 2, 3 и 4"}
    pname = period_names.get(n_periods, "всех")
    total_fk_days = last_dis + 1  # от Дня 0 до Дня last_dis включительно

    return (
        f"Максимальная продолжительность участия в исследовании для одного добровольца составит "
        f"{total} {_days_word(total)}. Период скрининга продлится от 1 до 14 дней (День от -14 до -1). "
        f"Длительность Периодов {pname} ФК – {total_fk_days} {_days_word(total_fk_days)} "
        f"(от Дня 0 до Дня {last_dis} включительно), "
        f"включая {wo} {_days_word(wo)} отмывочного периода между каждыми двумя последовательными периодами. "
        f"Период наблюдения – {fup} {_days_word(fup)} (День {last_dos} – День {fv} включительно) "
        f"от приёма последней дозы исследуемого препарата."
    )

# ─── Маппинг: строка → список значений для замены ────────────────────────────
# Каждое значение заменяет СЛЕДУЮЩЕЕ вхождение PLACEHOLDER в ячейке (слева направо, сверху вниз).
# Ключ — индекс строки таблицы (0-based).
# None — оставить placeholder как есть (не заменять; просто «пропустить» вхождение).
# Добавляй/изменяй нужные строки под свой JSON.

def build_field_map(args, data: dict, design_synopsis: str = "") -> dict:
    """Возвращает маппинг строк таблицы → список значений для замены."""
    fed_state    = args.fed_state
    dosing       = args.dosing

    # ── Хронология исследования ───────────────────────────────────────────────
    tl = _compute_study_timeline(data)
    dosing1      = DOSING_DESCRIPTIONS.get(dosing, ["однократного", "однократным"])[0]
    dosing2      = DOSING_DESCRIPTIONS.get(dosing, ["однократного", "однократным"])[1]

    drug         = data.get("drug", "")
    dosage_form  = data.get("dosage_form", "") or ""
    n_subjects   = str(data.get("n_subjects", ""))
    cv_intra     = str(data.get("cv_intra_used", ""))
    design       = data.get("design", "")

    ss           = data.get("sample_size_calculation", {})
    ss_reasoning = ss.get("reasoning", "")
    n_total      = str(ss.get("n_total", n_subjects))

    drug_cap     = drug.capitalize() if drug else ""

    # Дозировка из CLI-параметров (опционально)
    strength     = getattr(args, "strength", None)
    strength_unit = getattr(args, "strength_unit", "mg") or "mg"
    dose_number  = getattr(args, "dose_number", None)
    dose_unit    = getattr(args, "dose_unit", "mg") or "mg"

    # Форматированная строка дозировки, если указана
    def _fmt_strength():
        if strength is None:
            return None
        s = int(strength) if strength == int(strength) else strength
        return f"{s} {strength_unit}"

    def _fmt_dose():
        if dose_number is None:
            return None
        d = int(dose_number) if dose_number == int(dose_number) else dose_number
        return f"{d} {dose_unit}"

    strength_str = _fmt_strength()
    dose_str     = _fmt_dose()

    # Название препарата для строки 6 (без торгового наименования)
    drug6_parts = ["<Название тестового препарата>", f"ИНН: {drug_cap}"]
    if dosage_form:
        drug6_parts.append(dosage_form)
    if strength_str:
        drug6_parts.append(strength_str)
    drug6 = ", ".join(drug6_parts)

    # Заголовок исследования (строка 0)
    design_adj = DESIGN_TITLE_DESCRIPTIONS.get(design, "рандомизированном открытом")
    fed_state_title = {
        "натощак":                        "натощак",
        "после приема высококалорийной пищи": "после приёма высококалорийной пищи",
        "оба варианта":                   "натощак и после приёма пищи",
    }.get(fed_state, fed_state)

    study_title = (
        f"Открытое рандомизированное {design_adj} исследование сравнительной "
        f"фармакокинетики (биоэквивалентности) препаратов "
        f"<Название тестового препарата> (ИНН: {drug_cap}"
        + (f", {dosage_form}" if dosage_form else "")
        + (f", {strength_str}" if strength_str else "")
        + f") и <Название референтного препарата> (ИНН: {drug_cap}"
        + (f", {dosage_form}" if dosage_form else "")
        + (f", {strength_str}" if strength_str else "")
        + f") {fed_state_title} у здоровых добровольцев."
    )

    return {
        # ── Строка 0: Название клинического исследования ─────────────────────
        0: ["__REPLACE__", study_title],

        # ── Строки 1–4: идентификаторы / центры — заполни вручную ────────────
        # 1: Номер исследования
        # 2: Спонсор исследования
        # 3: Исследовательский центр
        # 4: Биоаналитическая лаборатория

        # ── Строка 5: Фаза — в шаблоне уже заполнена ─────────────────────────

        # ── Строка 6: Название исследуемого препарата ────────────────────────
        6: [drug6],

        # ── Строка 7: Действующее вещество ───────────────────────────────────
        7: [drug],

        # ── Строка 8: Цель исследования ──────────────────────────────────────
        8: [
            "__REPLACE__", _render("study_goal1", Drug=drug_cap, drug=drug, fed_state=fed_state),
            "__APPEND__",  _render("study_goal2", Drug=drug_cap, drug=drug, dosing_description=dosing1),
        ],

        # ── Строка 9: Задачи исследования ────────────────────────────────────
        9: ["__REPLACE__", _render("study_tasks", Drug=drug_cap, drug=drug, dosing1=dosing1, dosing2=dosing2)],

        # ── Строка 10: Дизайн исследования — LLM-генерированный текст ────────
        10: ["__REPLACE__", design_synopsis],

        # ── Строка 11: Методология ───────────────────────────────────────────
        11: ["__REPLACE__", _generate_methodology_text(tl, drug, dosing, fed_state)],

        # ── Строка 12: Количество добровольцев ───────────────────────────────
        12: ["__REPLACE__", _render(
            "sample_size_justification",
            cv_intra=cv_intra, drug=drug,
            n_total=n_total, n_subjects=n_subjects,
        )],

        # ── Строка 13: Критерии включения ────────────────────────────────────
        13: ["__REPLACE__", INCLUSION_CRITERIA],

        # ── Строка 14: Критерии невключения ──────────────────────────────────
        14: ["__REPLACE__", NON_INCLUSION_CRITERIA.format(drug=drug)],

        # ── Строка 15: Критерии исключения ───────────────────────────────────
        15: ["__REPLACE__", EXCLUSION_CRITERIA.format(drug=drug)],

        # ── Строка 16: Исследуемый препарат (T) ──────────────────────────────
        # Первый placeholder — торговое название (неизвестно → плейсхолдер)
        16: [
            "<Название тестового препарата>",  # торговое название T
            dosage_form,                        # лекарственная форма
            strength_str,                       # дозировка (None → оставит placeholder)
            dosage_form,                        # «Состав на 1 <лек.форма>»
            drug,                               # действующее вещество
            None,                               # вспомогательные вещества
            dose_str,                           # «по ___» (доза)
            drug_cap,                           # «препарата ___»
            None,                               # «в День ___»
            None, None, None,                   # оставшиеся
        ],

        # ── Строка 17: Референтный препарат (R) ──────────────────────────────
        17: [
            "<Название референтного препарата>",  # торговое название R
            drug,                                  # ИНН
            dosage_form,                           # лекарственная форма
            strength_str,                          # дозировка
            dosage_form,                           # «Состав на 1 ___»
            drug,                                  # действующее вещество
            None,                                  # вспомогательные вещества
            dose_str,                              # «по ___»
            drug_cap,                              # «препарата ___»
            None, None, None, None, None, None, None,
        ],

        # ── Строка 18: Периоды исследования ──────────────────────────────────
        18: ["__REPLACE__", _generate_periods_text(tl, drug, dosing, fed_state)],

        # ── Строка 19: Продолжительность исследования ─────────────────────────
        19: ["__REPLACE__", _generate_duration_text(tl)],

        # ── Строка 20: ФК параметры ───────────────────────────────────────────
        20: ["__REPLACE__", _render("pk_parameters", drug=drug)],

        # ── Строка 21: Аналитический метод ───────────────────────────────────
        21: [drug],

        # ── Строка 22: Критерии биоэквивалентности ────────────────────────────
        22: ["__REPLACE__", _render("be_criteria", Drug=drug_cap)],

        # ── Строка 23: Анализ безопасности ───────────────────────────────────
        23: [drug_cap],

        # ── Строка 24: Расчёт размера выборки (детальный — из JSON) ──────────
        24: ["__APPEND__", ss_reasoning],

        # ── Строка 25: Методы статистического анализа ────────────────────────
        25: ["__REPLACE__", _render("statistical_methods", Drug=drug_cap)],

        # ── Строка 26: Заслепление и Рандомизация ────────────────────────────
        26: [drug_cap],

        # ── Строка 27: Этические и регуляторные аспекты ──────────────────────
        27: [drug_cap],

        # ── Строка 28: Номер версии протокола ────────────────────────────────
        # 28: [None],
    }


# ─── Утилиты для работы с docx ────────────────────────────────────────────────

def _replace_in_paragraph_once(paragraph, old: str, new: str) -> bool:
    """
    Заменяет первое вхождение `old` на `new` в параграфе.
    Сначала пробует покраснить в отдельных runs, затем — пересборкой.
    Возвращает True, если замена произошла.
    """
    # Быстрая проверка по полному тексту
    full = "".join(r.text for r in paragraph.runs)
    if old not in full:
        return False

    # Попытка 1: замена внутри одного run (без потери форматирования)
    for run in paragraph.runs:
        if old in run.text:
            run.text = run.text.replace(old, new, 1)
            return True

    # Попытка 2: placeholder разбит по нескольким runs — собираем в первый
    if paragraph.runs:
        paragraph.runs[0].text = full.replace(old, new, 1)
        for run in paragraph.runs[1:]:
            run.text = ""
    return True

def _is_bold_header(line: str) -> bool:
    # нормализация: NBSP, thin space, em space → обычный пробел, затем strip
    s = (line or "").replace("\u00A0", " ").replace("\u2009", " ").replace("\u2002", " ").strip().lower()
    return s in {"основная цель:", "дополнительная цель:", "дополнительная цель :"} \
        or s.startswith("основная цель:") \
        or s.startswith("дополнительная цель:")

def _append_text_to_cell(cell, text: str):
    """Добавляет текст в конец ячейки и делает жирными строки-заголовки."""
    lines = str(text).split("\n")
    for line in lines:
        if line.strip() == "":
            continue

        p = cell.add_paragraph()
        _set_para_compact(p)

        run = p.add_run(line)

        if _is_bold_header(line):
            run.bold = True

# def _append_text_to_cell(cell, text: str):
#     """Добавляет текст в пустую ячейку, разбивая по абзацам."""
#     lines = text.split("\n")
#     if not lines:
#         return

#     # Первый параграф уже существует в docx-ячейке, используем его
#     first_para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
#     if first_para.runs:
#         first_para.runs[0].text = lines[0]
#     else:
#         first_para.add_run(lines[0])

#     base_style = first_para.style
#     for line in lines[1:]:
#         p = cell.add_paragraph(line)
#         p.style = base_style

def _set_para_compact(p):
    """Убираем большие интервалы между строками/абзацами."""
    pf = p.paragraph_format
    pf.space_before = 0
    pf.space_after = 0
    pf.line_spacing = 1.0  # можно 1.15 если нужно

def _replace_cell_content(cell, text: str):
    """
    Полностью заменяет содержимое ячейки, создавая абзацы заново.
    Делает жирным строки-заголовки.
    """
    lines = str(text).split("\n")

    # ВАЖНО: это удаляет ВСЕ старые абзацы ячейки (и их "пустые" интервалы)
    cell.text = ""

    for i, line in enumerate(lines):
        # пропускаем полностью пустые строки, чтобы не было лишних разрывов
        if line.strip() == "":
            continue

        p = cell.paragraphs[0] if (len(cell.paragraphs) == 1 and cell.paragraphs[0].text == "" and i == 0) else cell.add_paragraph()
        _set_para_compact(p)

        run = p.add_run(line)

        # Жирные заголовки
        if line.strip().startswith("Основная цель:") or line.strip().startswith("Дополнительная цель:"):
            run.bold = True
            
# def _replace_cell_content(cell, text: str):
#     """
#     Заменяет ВСЁ содержимое ячейки новым текстом (разбивая по «\\n»).
#     Используется с маркером __REPLACE__ — когда шаблонный текст генерируется
#     в коде, а не подбирается из заглушек docx.
#     """
#     lines = text.split("\n")

#     # Очищаем все существующие runs (сохраняем параграфы и стиль)
#     for para in cell.paragraphs:
#         for run in para.runs:
#             run.text = ""

#     first_para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
#     if first_para.runs:
#         first_para.runs[0].text = lines[0]
#     else:
#         first_para.add_run(lines[0])

#     base_style = first_para.style
#     for line in lines[1:]:
#         p = cell.add_paragraph(line)
#         p.style = base_style


def fill_cell(cell, values: list):
    """
    Заменяет вхождения PLACEHOLDER в ячейке последовательно значениями из `values`.
    - None           → оставить placeholder без изменений, но «потребить» позицию
    - "__APPEND__"   → следующее значение добавляется в пустую ячейку целиком
    - "__REPLACE__"  → следующее значение заменяет ВСЁ содержимое ячейки
    """
    # Временный маркер для позиций, которые нужно оставить как placeholder
    SENTINEL = "\ufeff" * 5

    append_mode = False
    replace_mode = False
    for val in values:
        if val == "__APPEND__":
            append_mode = True
            continue
        if val == "__REPLACE__":
            replace_mode = True
            continue

        if append_mode:
            _append_text_to_cell(cell, str(val))
            append_mode = False
            continue

        if replace_mode:
            _replace_cell_content(cell, str(val))
            replace_mode = False
            continue

        if val is None:
            # Потребляем позицию: временно заменяем на sentinel, чтобы следующий
            # вызов нашёл уже СЛЕДУЮЩЕЕ вхождение PLACEHOLDER
            for para in cell.paragraphs:
                if _replace_in_paragraph_once(para, PLACEHOLDER, SENTINEL):
                    break
            continue

        # Заменяем первый оставшийся PLACEHOLDER на значение
        for para in cell.paragraphs:
            if _replace_in_paragraph_once(para, PLACEHOLDER, str(val)):
                break

    # Восстанавливаем sentinel → исходный placeholder
    for para in cell.paragraphs:
        for run in para.runs:
            if SENTINEL in run.text:
                run.text = run.text.replace(SENTINEL, PLACEHOLDER)


# ─── Основная функция ─────────────────────────────────────────────────────────

def _find_bibliography_row(table) -> int:
    """
    Находит строку таблицы, в первой ячейке которой содержится 'библиограф'.
    Возвращает индекс строки или -1 если не найдено.
    """
    for i, row in enumerate(table.rows):
        if row.cells:
            cell_text = "".join(p.text for p in row.cells[0].paragraphs).lower()
            if "библиограф" in cell_text:
                return i
    return -1


def _format_references(records) -> str:
    """
    Форматирует библиографический список из списка PKRecord.

    Для PubMed: название + PMID + ссылка на pubmed.ncbi.nlm.nih.gov
    Для OpenFDA: название + номер заявки FDA
    """
    entries = []
    n = 0
    for r in records:
        if r.source == "merged":
            continue  # синтетическая запись, не реальный источник
        if not r.title and not r.study_id:
            continue
        n += 1
        title = r.title or r.source
        if r.source == "PubMed" and r.study_id:
            pmid = r.study_id
            entries.append(
                f"{n}. {title}\n"
                f"   PMID: {pmid}. https://pubmed.ncbi.nlm.nih.gov/{pmid}/"
            )
        elif r.source == "OpenFDA":
            app_line = f"\n   FDA Application No.: {r.study_id}" if r.study_id else ""
            entries.append(f"{n}. {title}{app_line}")
        else:
            id_part = f" (ID: {r.study_id})" if r.study_id else ""
            entries.append(f"{n}. {title}{id_part}")
    return "\n\n".join(entries)


def fill_template(
    args: argparse.Namespace,
    json_path: str,
    output_path: str,
    template_path: str = "",
    design_synopsis: str = "",
    records=None,
) -> None:
    """
    Заполняет шаблон данными из JSON и сохраняет результат.

    Args:
        json_path:       путь к JSON-файлу с данными
        output_path:     путь для сохранения заполненного docx
        template_path:   путь к шаблону
        design_synopsis: LLM-генерированное обоснование дизайна для строки 10
        records:         список PKRecord — источники данных для библиографии
    """
    with open(json_path, encoding="utf-8") as f:
        data = json.load(f)

    doc = Document(template_path)
    table = doc.tables[0]

    field_map = build_field_map(args, data, design_synopsis)

    for row_idx, values in field_map.items():
        if row_idx >= len(table.rows):
            print(f"  [!] Строка {row_idx} не существует в таблице — пропускаю.")
            continue
        row = table.rows[row_idx]
        if len(row.cells) < 2:
            print(f"  [!] Строка {row_idx} не имеет второй ячейки — пропускаю.")
            continue
        fill_cell(row.cells[1], values)

    # ── Библиографический список источников ──────────────────────────────────
    if records:
        refs_text = _format_references(records)
        if refs_text:
            bib_row = _find_bibliography_row(table)
            if bib_row >= 0 and len(table.rows[bib_row].cells) >= 2:
                _replace_cell_content(table.rows[bib_row].cells[1], refs_text)
                print(f"  ✓ Библиография: {refs_text.count(chr(10) + chr(10)) + 1} источника(ов) добавлено в строку {bib_row}")
            else:
                print("  [!] Строка 'Библиографический список' не найдена в шаблоне — пропускаю.")

    doc.save(output_path)
    print(f"✓ Синопсис сохранён: {output_path}")


# ─── Точка входа ──────────────────────────────────────────────────────────────

# if __name__ == "__main__":
    # json_file = sys.argv[1] if len(sys.argv) > 1 else "recommendation_amlodipine.json"
    # out_file  = sys.argv[2] if len(sys.argv) > 2 else "synopsis_output.docx"

    # if not Path(json_file).exists():
    #     print(f"Файл не найден: {json_file}")
    #     sys.exit(1)
    # if not Path(template_path := TEMPLATE_PATH).exists():
    #     print(f"Шаблон не найден: {template_path}")
    #     sys.exit(1)

    # fill_template(json_file, out_file)
